#!/usr/bin/env python3
"""
Ingestion script for swim plan templates.
Reads .docx files from source folders and outputs JSON to data/templates.v1.json
"""

import os
import json
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("Error: python-docx is not installed.")
    print("Please install it with: pip install python-docx")
    sys.exit(1)


# Template type configuration
TEMPLATE_TYPES = [
    {
        "plan_type_key": "mileage",
        "plan_type_label": "Mileage Mondays",
        "source_folder": "swim_templates/source/mileage"
    },
    {
        "plan_type_key": "im",
        "plan_type_label": "IM Sessions",
        "source_folder": "swim_templates/source/im"
    },
    {
        "plan_type_key": "fast",
        "plan_type_label": "Fast Sessions",
        "source_folder": "swim_templates/source/fast"
    },
    {
        "plan_type_key": "kitchen_sink",
        "plan_type_label": "Kitchen Sink Sessions",
        "source_folder": "swim_templates/source/kitchen_sink"
    }
]

OUTPUT_FILE = "data/templates.v1.json"


def extract_text_from_docx(docx_path):
    """
    Extract all text from a .docx file including paragraphs and tables.

    Args:
        docx_path: Path to the .docx file

    Returns:
        str: Extracted text content
    """
    doc = Document(docx_path)
    extracted_text = []

    # Extract paragraph text
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            extracted_text.append(text)

    # Extract table text
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                extracted_text.append(" | ".join(row_text))

    return "\n".join(extracted_text)


def find_first_docx(folder_path):
    """
    Find the first .docx file in a folder.

    Args:
        folder_path: Path to the folder to search

    Returns:
        Path object or None if no .docx file found
    """
    folder = Path(folder_path)
    if not folder.exists():
        return None

    for file in folder.iterdir():
        if file.suffix.lower() == '.docx' and not file.name.startswith('~'):
            return file

    return None


def main():
    """Main ingestion process."""
    script_dir = Path(__file__).parent.parent.parent  # Go up to repo root
    os.chdir(script_dir)

    print("🏊 Swim Template Ingestion v1")
    print("=" * 50)

    templates = []
    missing_templates = []

    for template_type in TEMPLATE_TYPES:
        folder = template_type["source_folder"]
        key = template_type["plan_type_key"]
        label = template_type["plan_type_label"]

        print(f"\n📄 Processing {label} ({key})...")

        docx_file = find_first_docx(folder)

        if docx_file is None:
            print(f"   ⚠️  No .docx file found in {folder}/")
            missing_templates.append(key)
            continue

        print(f"   Found: {docx_file.name}")

        try:
            raw_text = extract_text_from_docx(docx_file)
            char_count = len(raw_text)
            line_count = len(raw_text.split('\n'))

            print(f"   ✓ Extracted {char_count} characters, {line_count} lines")

            templates.append({
                "plan_type_key": key,
                "plan_type_label": label,
                "source_file": docx_file.name,
                "raw_text": raw_text
            })
        except Exception as e:
            print(f"   ✗ Error extracting text: {e}")
            missing_templates.append(key)

    # Write output
    print(f"\n📝 Writing output to {OUTPUT_FILE}...")

    output_data = {
        "templates": templates,
        "version": "1.0",
        "generated_at": None  # Could add timestamp if needed
    }

    output_path = Path(OUTPUT_FILE)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(output_data, f, indent=2, ensure_ascii=False)

    print(f"✓ Successfully wrote {len(templates)} templates to {OUTPUT_FILE}")

    if missing_templates:
        print(f"\n⚠️  Warning: Missing templates for: {', '.join(missing_templates)}")
        print("   The AI coach will work with partial templates.")

    print("\n✅ Ingestion complete!")
    print(f"   Total templates: {len(templates)}/4")

    return 0 if len(templates) > 0 else 1


if __name__ == "__main__":
    sys.exit(main())
