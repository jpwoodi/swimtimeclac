#!/usr/bin/env python3
"""
Enhanced ingestion script for swim plan templates (v2).
Processes bulk .docx files from source folders and outputs JSON with rich metadata.
"""

import os
import json
import sys
import re
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
except ImportError:
    print("Error: python-docx is not installed.")
    print("Please install it with: pip install python-docx")
    sys.exit(1)


# Template type configuration
TEMPLATE_TYPES = [
    {
        "plan_type_key": "mileage",
        "plan_type_label": "Mileage (Distance)",
        "source_folder": "swim_templates/source/mileage"
    },
    {
        "plan_type_key": "im",
        "plan_type_label": "IM (Strokes)",
        "source_folder": "swim_templates/source/im"
    },
    {
        "plan_type_key": "fast",
        "plan_type_label": "Fast (Speed)",
        "source_folder": "swim_templates/source/fast"
    },
    {
        "plan_type_key": "kitchen_sink",
        "plan_type_label": "Kitchen Sink (Mixed)",
        "source_folder": "swim_templates/source/kitchen_sink"
    }
]

OUTPUT_FILE = "data/templates.v2.json"
MOJIBAKE_REPLACEMENTS = {
    "\u2192": "->",  # Right arrow
    "\u2013": "-",   # En dash
    "\u2014": "-",   # Em dash
    "\u2018": "'",   # Left single quote
    "\u2019": "'",   # Right single quote
    "\u201c": '"',   # Left double quote
    "\u201d": '"',   # Right double quote
    "\u00a0": " "    # Non-breaking space
}


def normalize_text(text):
    """Normalize extracted text and repair common mojibake artifacts."""
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    for bad, good in MOJIBAKE_REPLACEMENTS.items():
        normalized = normalized.replace(bad, good)
    return normalized.strip()


def extract_text_from_docx(docx_path):
    """
    Extract all text from a .docx file including paragraphs and tables.

    Args:
        docx_path: Path to the .docx file

    Returns:
        str: Extracted text content
    """
    doc = Document(docx_path)
    extracted_text = []

    # Extract paragraph text
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            extracted_text.append(text)

    # Extract table text
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                extracted_text.append("\t".join(row_text))

    return normalize_text("\n".join(extracted_text))


def extract_distance_from_content(raw_text):
    """
    Extract total distance from plan content.
    Looks for patterns like "3,200", "2800", "3000 yards", etc.
    """
    # Look for distance at end of text (common pattern: Cool Down\t3,200)
    distance_patterns = [
        r'Cool Down.*?(\d{1,2}[,.]?\d{3})',  # After "Cool Down"
        r'Total.*?(\d{1,2}[,.]?\d{3})',      # After "Total"
        r'(\d{1,2}[,.]?\d{3})\s*(?:yards|meters|m|y)',  # Explicit units
    ]

    for pattern in distance_patterns:
        match = re.search(pattern, raw_text, re.IGNORECASE)
        if match:
            distance_str = match.group(1).replace(',', '').replace('.', '')
            try:
                return int(distance_str)
            except ValueError:
                pass

    return None


def extract_metadata_from_filename(filename):
    """
    Extract metadata from filename patterns.
    Expected format: "YYYY.MM.DD - DISTANCE - POOL_TYPE.docx"
    Example: "2026.02.09 - 3200 - SCY.docx"
    """
    metadata = {
        "date": None,
        "distance_meters": None,
        "pool_type": None
    }

    # Remove .docx extension
    name = filename.replace('.docx', '').replace('.DOCX', '')

    # Extract date (YYYY.MM.DD or YYYY-MM-DD)
    date_match = re.match(r'^(\d{4})[.\-](\d{2})[.\-](\d{2})', name)
    if date_match:
        metadata["date"] = f"{date_match.group(1)}-{date_match.group(2)}-{date_match.group(3)}"
        # Remove the date part to avoid false distance matching
        name = name[date_match.end():]

    # Extract distance (e.g., "3200", "2800")
    # Look for 4-digit number NOT in date position
    distance_match = re.search(r'\b(\d{4})\b', name)
    if distance_match:
        metadata["distance_meters"] = int(distance_match.group(1))

    # Extract pool type (SCY, SCM, LCM)
    pool_match = re.search(r'\b(SCY|SCM|LCM)\b', name, re.IGNORECASE)
    if pool_match:
        metadata["pool_type"] = pool_match.group(1).upper()

    return metadata


def classify_difficulty(distance_meters):
    """
    Classify difficulty based on total distance.
    This is a basic heuristic and can be refined.
    """
    if distance_meters is None:
        return "intermediate"

    if distance_meters < 2000:
        return "beginner"
    elif distance_meters < 3000:
        return "intermediate"
    elif distance_meters < 4000:
        return "advanced"
    else:
        return "elite"


def extract_equipment_mentioned(raw_text):
    """
    Extract equipment mentioned in the plan text.
    """
    equipment_keywords = {
        "pull_buoy": ["pull buoy", "pull", "pulling"],
        "fins": ["fins", "fin"],
        "kickboard": ["kick board", "kickboard", "kick"],
        "paddles": ["paddles", "paddle"],
        "snorkel": ["snorkel"],
        "band": ["band", "ankle band"]
    }

    equipment_found = []
    raw_text_lower = raw_text.lower()

    for equipment, keywords in equipment_keywords.items():
        if any(keyword in raw_text_lower for keyword in keywords):
            equipment_found.append(equipment)

    return equipment_found


def extract_focus_areas(raw_text, plan_type_key):
    """
    Extract focus areas from plan content and type.
    """
    focus_areas = []
    raw_text_lower = raw_text.lower()

    # Type-based focus
    type_focus = {
        "mileage": ["endurance", "volume"],
        "im": ["technique", "stroke_work"],
        "fast": ["speed", "sprint"],
        "kitchen_sink": ["mixed", "variety"]
    }

    focus_areas.extend(type_focus.get(plan_type_key, []))

    # Content-based focus
    if any(word in raw_text_lower for word in ["sprint", "fast", "speed"]):
        if "speed" not in focus_areas:
            focus_areas.append("speed")

    if any(word in raw_text_lower for word in ["drill", "technique"]):
        if "technique" not in focus_areas:
            focus_areas.append("technique")

    if any(word in raw_text_lower for word in ["kick", "kicking"]):
        if "kick" not in focus_areas:
            focus_areas.append("kick")

    return focus_areas


def estimate_duration(distance_meters):
    """
    Estimate session duration based on distance.
    Assumes average pace of ~1:45/100m including rest.
    """
    if distance_meters is None:
        return 60

    # Conservative estimate: 1.75 min per 100m (includes rest)
    duration = (distance_meters / 100) * 1.75

    # Round to nearest 5 minutes
    return int(round(duration / 5) * 5)


def find_all_docx(folder_path):
    """
    Find all .docx files in a folder.

    Args:
        folder_path: Path to the folder to search

    Returns:
        List of Path objects
    """
    folder = Path(folder_path)
    if not folder.exists():
        return []

    docx_files = []
    for file in folder.iterdir():
        if file.suffix.lower() == '.docx' and not file.name.startswith('~') and not file.name.startswith('.'):
            docx_files.append(file)

    # Sort by name for consistent ordering
    return sorted(docx_files, key=lambda x: x.name)


def generate_plan_id(plan_type_key, filename):
    """
    Generate a unique plan ID from type and filename.
    Example: "mileage_2026-02-09_3200_scy"
    """
    # Remove extension and special characters
    clean_name = filename.replace('.docx', '').replace('.DOCX', '')
    clean_name = re.sub(r'[^a-zA-Z0-9\-_.]', '_', clean_name)

    return f"{plan_type_key}_{clean_name}".lower()


def process_template(docx_file, template_type):
    """
    Process a single template file and extract all metadata.
    """
    try:
        raw_text = extract_text_from_docx(docx_file)
        filename_meta = extract_metadata_from_filename(docx_file.name)

        # Extract distance (prefer from filename, fallback to content)
        distance_meters = filename_meta.get("distance_meters")
        if distance_meters is None:
            distance_meters = extract_distance_from_content(raw_text)

        # Build complete template object
        template = {
            "plan_id": generate_plan_id(template_type["plan_type_key"], docx_file.name),
            "plan_type_key": template_type["plan_type_key"],
            "plan_type_label": template_type["plan_type_label"],
            "source_file": docx_file.name,
            "raw_text": raw_text,
            "metadata": {
                "date": filename_meta.get("date"),
                "distance_meters": distance_meters,
                "pool_type": filename_meta.get("pool_type", "SCY"),
                "difficulty": classify_difficulty(distance_meters),
                "focus_areas": extract_focus_areas(raw_text, template_type["plan_type_key"]),
                "equipment_required": extract_equipment_mentioned(raw_text),
                "estimated_duration_minutes": estimate_duration(distance_meters),
                "intensity": "medium"  # Default, can be enhanced later
            }
        }

        return template

    except Exception as e:
        print(f"      ERROR: {e}")
        return None


def main():
    """Main ingestion process for v2."""
    script_dir = Path(__file__).parent.parent.parent  # Go up to repo root
    os.chdir(script_dir)

    print("Swim Template Ingestion v2 (Bulk Processing)")
    print("=" * 60)

    all_templates = []
    stats = {
        "total_files": 0,
        "successful": 0,
        "failed": 0,
        "by_type": {}
    }

    for template_type in TEMPLATE_TYPES:
        folder = template_type["source_folder"]
        key = template_type["plan_type_key"]
        label = template_type["plan_type_label"]

        print(f"\n[{key.upper()}] Processing {label}...")
        print(f"   Folder: {folder}")

        docx_files = find_all_docx(folder)

        if not docx_files:
            print(f"   WARNING: No .docx files found")
            stats["by_type"][key] = 0
            continue

        print(f"   Found {len(docx_files)} files")

        type_count = 0
        for docx_file in docx_files:
            stats["total_files"] += 1
            print(f"      • {docx_file.name}", end=" ... ")

            template = process_template(docx_file, template_type)

            if template:
                all_templates.append(template)
                stats["successful"] += 1
                type_count += 1

                # Show key metadata
                distance = template["metadata"]["distance_meters"]
                difficulty = template["metadata"]["difficulty"]
                print(f"OK ({distance}m, {difficulty})")
            else:
                stats["failed"] += 1
                print("FAILED")

        stats["by_type"][key] = type_count
        print(f"   Processed: {type_count} templates")

    # Write output
    print(f"\n{'=' * 60}")
    print(f"Writing output to {OUTPUT_FILE}...")

    output_data = {
        "templates": all_templates,
        "version": "2.0",
        "generated_at": datetime.utcnow().isoformat() + "Z",
        "stats": stats
    }

    output_path = Path(OUTPUT_FILE)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(output_data, f, indent=2, ensure_ascii=False)

    print(f"✓ Successfully wrote {len(all_templates)} templates")
    print(f"\nStatistics:")
    print(f"   Total files processed: {stats['total_files']}")
    print(f"   Successful: {stats['successful']}")
    print(f"   Failed: {stats['failed']}")
    print(f"\nBy type:")
    for key, count in stats["by_type"].items():
        print(f"   {key}: {count} templates")

    # Show difficulty distribution
    difficulty_dist = {}
    for template in all_templates:
        diff = template["metadata"]["difficulty"]
        difficulty_dist[diff] = difficulty_dist.get(diff, 0) + 1

    print(f"\nDifficulty distribution:")
    for diff, count in sorted(difficulty_dist.items()):
        print(f"   {diff}: {count}")

    print(f"\n{'=' * 60}")
    print("Ingestion complete!")

    return 0 if stats["successful"] > 0 else 1


if __name__ == "__main__":
    sys.exit(main())
